import streamlit as st
import pandas as pd
import os
from io import BytesIO
from docx import Document as DocxDocument
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime

# --- CẤU HÌNH ---
st.set_page_config(page_title="Quản lý dữ liệu Công trình", layout="wide")

DB_FILE = 'database_cong_trinh.xlsx'

ALL_COLUMNS = [
    'STT', 'Tên Công trình', 'Mã CT', 'Kế hoạch', 'Số Phương án', 'Ngày Phương án', 
    'Giá trị Phương án', 'Số Dự toán', 'Ngày Dự toán', 'Giá trị Dự toán', 
    'Số Hợp đồng thiết kế', 'Ngày Hợp đồng thiết kế', 'Giá trị Hợp đồng thiết kế', 
    'Số Hợp đồng giám sát', 'Ngày Hợp đồng giám sát', 'Giá trị Hợp đồng giám sát', 
    'Số Hợp đồng xây lắp', 'Ngày Hợp đồng xây lắp', 'Giá trị Hợp đồng xây lắp', 
    'Giá trị phát sinh', 'Giá trị VT thừa', 'Giá trị VTTH', 
    'Số Q.định phê duyệt QT công trình', 'Ngày Q.định phê duyệt QT công trình', 
    'Giá trị Q.định phê duyệt QT công trình', 'Số tiền bằng chữ', 'Ghi chú', 'Đơn vị QL',
    'Căn cứ pháp lý', 'Khối lượng công việc', 'Ngày khởi công', 'Ngày hoàn thành'
]

# Hàm load/khởi tạo dữ liệu
def load_data():
    if os.path.exists(DB_FILE):
        return pd.read_excel(DB_FILE)
    else:
        # Tạo file mới với các cột chuẩn
        df = pd.DataFrame(columns=ALL_COLUMNS)
        df.to_excel(DB_FILE, index=False)
        return df

def doc_so_vn(n):
    if not n: return "Không đồng"
    try:
        n = int(n)
    except:
        return ""
    if n == 0: return "Không đồng"
    if n < 0: return "Âm " + doc_so_vn(-n)
    
    digits = ["không", "một", "hai", "ba", "bốn", "năm", "sáu", "bảy", "tám", "chín"]
    units = ["", "nghìn", "triệu", "tỷ"]
    
    def decode_3(num, full):
        res = []
        h = num // 100
        t = (num % 100) // 10
        u = num % 10
        if full or h > 0:
            res.append(f"{digits[h]} trăm")
        if t > 1:
            res.append(f"{digits[t]} mươi")
            if u == 1: res.append("mốt")
            elif u == 5: res.append("lăm")
            elif u > 0: res.append(digits[u])
        elif t == 1:
            res.append("mười")
            if u == 5: res.append("lăm")
            elif u > 0: res.append(digits[u])
        elif res and u > 0:
            res.append(f"lẻ {digits[u]}")
        elif not res and u > 0:
            res.append(digits[u])
        return " ".join(res)

    blocks = []
    while n > 0:
        blocks.append(n % 1000)
        n //= 1000

    words = []
    for i, block in enumerate(blocks):
        if block == 0 and i > 0: continue
        full = (i < len(blocks)-1)
        unit = units[i % 4]
        for _ in range(i // 4): unit += " tỷ" 
        
        words.append(f"{decode_3(block, full)} {unit}".strip())
    
    res = " ".join(reversed(words)).strip()
    res = res.replace("  ", " ").strip()
    return res[0].upper() + res[1:] + " đồng"

def format_num_val(v):
    if pd.isna(v) or v == "" or v is None: return "0"
    try:
        return f"{int(float(v)):,}"
    except:
        return "0"

def parse_num_val(s):
    if not s: return 0
    try:
        return int(str(s).replace(',', '').replace('.', '').strip())
    except:
        return 0

# --- GIAO DIỆN ---
st.title("Phần mềm Nhập liệu và Báo cáo Công trình")

tab1, tab2, tab3 = st.tabs(["📝 Nhập liệu", "📊 Báo cáo", "📄 Bảng thuyết minh quyết toán"])

with tab1:
    st.header("1. Nhập thông tin Công trình chính")
    
    db_df_tab1 = load_data()
    main_mask_tab1 = db_df_tab1['Kế hoạch'].notna()
    list_cong_trinh_tab1 = db_df_tab1.loc[main_mask_tab1, 'Tên Công trình'].dropna().unique().tolist()
    
    selected_edit_ct = st.selectbox("Chọn Công trình để chỉnh sửa (hoặc Thêm mới)", ["-- Thêm mới --"] + list_cong_trinh_tab1)
    
    defaults = {
        'STT': "I", 'Tên Công trình': "", 'Mã CT': "", 'Kế hoạch': 0, 'Số Phương án': "", 'Ngày Phương án': None,
        'Giá trị Phương án': 0, 'Số Dự toán': "", 'Ngày Dự toán': None, 'Giá trị Dự toán': 0, 'Số Hợp đồng thiết kế': "",
        'Ngày Hợp đồng thiết kế': None, 'Giá trị Hợp đồng thiết kế': 0, 'Số Hợp đồng giám sát': "", 'Ngày Hợp đồng giám sát': None,
        'Giá trị Hợp đồng giám sát': 0, 'Số Hợp đồng xây lắp': "", 'Ngày Hợp đồng xây lắp': None, 'Giá trị Hợp đồng xây lắp': 0,
        'Giá trị phát sinh': 0, 'Giá trị VT thừa': 0, 'Giá trị VTTH': 0, 'Số Q.định phê duyệt QT công trình': "",
        'Ngày Q.định phê duyệt QT công trình': None, 'Giá trị Q.định phê duyệt QT công trình': 0, 'Ghi chú': "", 'Đơn vị QL': "",
        'Căn cứ pháp lý': "", 'Khối lượng công việc': "", 'Ngày khởi công': None, 'Ngày hoàn thành': None
    }

    sub_items_columns = ['STT', 'Tên Hạng mục', 'Giá trị Dự toán', 'Giá trị quyết toán', 'Chênh lệch']
    
    if selected_edit_ct != "-- Thêm mới --":
        start_indices = db_df_tab1.index[db_df_tab1['Tên Công trình'] == selected_edit_ct].tolist()
        if start_indices:
            start_idx = start_indices[0]
            end_idx = len(db_df_tab1)
            for i in range(start_idx + 1, len(db_df_tab1)):
                val = str(db_df_tab1.at[i, 'STT']).strip().upper()
                if val in ['I', 'II', 'III', 'IV', 'V', 'VI', 'VII', 'VIII', 'IX', 'X']:
                    end_idx = i
                    break
            
            active_df = db_df_tab1.iloc[start_idx:end_idx]
            main_row_data = active_df.iloc[0].to_dict()
            for k in defaults.keys():
                if pd.notna(main_row_data.get(k)):
                    defaults[k] = main_row_data[k]
                    if isinstance(defaults[k], pd.Timestamp):
                        defaults[k] = defaults[k].date()
            
            if st.session_state.get('last_edit_ct') != selected_edit_ct:
                sub_rows = active_df.iloc[1:].copy()
                sub_rows = sub_rows[['STT', 'Tên Công trình', 'Giá trị Dự toán', 'Giá trị Q.định phê duyệt QT công trình', 'Ghi chú']]
                sub_rows = sub_rows.rename(columns={'Tên Công trình': 'Tên Hạng mục', 'Giá trị Q.định phê duyệt QT công trình': 'Giá trị quyết toán', 'Ghi chú': 'Chênh lệch'})
                sub_rows['Giá trị Dự toán'] = sub_rows['Giá trị Dự toán'].fillna(0).astype(int)
                sub_rows['Giá trị quyết toán'] = sub_rows['Giá trị quyết toán'].fillna(0).astype(int)
                for col in ['STT', 'Tên Hạng mục', 'Chênh lệch']:
                    sub_rows[col] = sub_rows[col].fillna("")
                    
                st.session_state.sub_df = sub_rows
                st.session_state.last_edit_ct = selected_edit_ct
    else:
        if st.session_state.get('last_edit_ct') != "-- Thêm mới --":
            if 'sub_df' in st.session_state:
                del st.session_state['sub_df']
            st.session_state.last_edit_ct = "-- Thêm mới --"

    wid_key = selected_edit_ct
    
    with st.expander("Thông tin cơ bản", expanded=True):
        col1, col2, col3 = st.columns(3)
        with col1:
            stt = st.text_input("STT (VD: I, II)", value=str(defaults['STT']), key=f"stt_{wid_key}")
            ten_ct = st.text_input("Tên Công trình (* Bắt buộc)", value=str(defaults['Tên Công trình']), key=f"ten_ct_{wid_key}")
            don_vi_ql = st.text_input("Đơn vị Quản lý", value=str(defaults['Đơn vị QL']), key=f"dvql_{wid_key}")
        with col2:
            ma_ct = st.text_input("Mã CT", value=str(defaults['Mã CT']), key=f"ma_ct_{wid_key}")
            ke_hoach_str = st.text_input("Kế hoạch (Giá trị)", value=format_num_val(defaults['Kế hoạch']), key=f"kehoach_{wid_key}")
            ke_hoach = parse_num_val(ke_hoach_str)
            so_tien_chu = st.text_input("Số tiền bằng chữ", value=doc_so_vn(ke_hoach), key=f"sotien_{wid_key}")
        with col3:
            ghi_chu = st.text_input("Ghi chú", value=str(defaults['Ghi chú']), key=f"ghichu_{wid_key}")
            col_ngay1, col_ngay2 = st.columns(2)
            with col_ngay1:
                ngay_khoi_cong = st.date_input("Ngày khởi công", value=defaults['Ngày khởi công'], format="DD/MM/YYYY", key=f"ngaykc_{wid_key}")
            with col_ngay2:
                ngay_hoan_thanh = st.date_input("Ngày hoàn thành", value=defaults['Ngày hoàn thành'], format="DD/MM/YYYY", key=f"ngayht_{wid_key}")
        
        can_cu_phap_ly = st.text_area("Căn cứ pháp lý", value=str(defaults['Căn cứ pháp lý']), height=150, key=f"ccpl_{wid_key}", help="Nhập các căn cứ pháp lý liên quan, có thể xuống hàng")
        khoi_luong_cv = st.text_area("Khối lượng công việc", value=str(defaults['Khối lượng công việc']), height=150, key=f"klcv_{wid_key}", help="Nhập các khối lượng công việc, có thể xuống hàng")
    
    with st.expander("Thông tin Phương án, Dự toán và Phê duyệt QT"):
        col4, col5, col6 = st.columns(3)
        with col4:
            st.markdown("**Phương án**")
            so_pa = st.text_input("Số Phương án", value=str(defaults['Số Phương án']), key=f"sopa_{wid_key}")
            ngay_pa = st.date_input("Ngày Phương án", value=defaults['Ngày Phương án'], format="DD/MM/YYYY", key=f"ngaypa_{wid_key}")
            gt_pa_str = st.text_input("Giá trị Phương án", value=format_num_val(defaults['Giá trị Phương án']), key=f"gtpa_{wid_key}")
            gt_pa = parse_num_val(gt_pa_str)
        with col5:
            st.markdown("**Dự toán**")
            so_dt = st.text_input("Số Dự toán", value=str(defaults['Số Dự toán']), key=f"sodt_{wid_key}")
            ngay_dt = st.date_input("Ngày Dự toán", value=defaults['Ngày Dự toán'], format="DD/MM/YYYY", key=f"ngaydt_{wid_key}")
            gt_dt_str = st.text_input("Giá trị Dự toán", value=format_num_val(defaults['Giá trị Dự toán']), key=f"gtdt_{wid_key}")
            gt_dt = parse_num_val(gt_dt_str)
        with col6:
            st.markdown("**QĐ Phê duyệt QT CT**")
            so_qd = st.text_input("Số Q.định phê duyệt QT công trình", value=str(defaults['Số Q.định phê duyệt QT công trình']), key=f"soqd_{wid_key}")
            ngay_qd = st.date_input("Ngày Q.định phê duyệt QT công trình", value=defaults['Ngày Q.định phê duyệt QT công trình'], format="DD/MM/YYYY", key=f"ngayqd_{wid_key}")
            gt_qd_str = st.text_input("Giá trị Q.định phê duyệt QT công trình", value=format_num_val(defaults['Giá trị Q.định phê duyệt QT công trình']), key=f"gtqd_{wid_key}")
            gt_qd = parse_num_val(gt_qd_str)
            
    with st.expander("Thông tin Hợp đồng & Vật tư khác (Tùy chọn)"):
        c1, c2, c3 = st.columns(3)
        with c1:
            st.markdown("**HĐ Thiết kế**")
            so_hdtk = st.text_input("Số HĐ TK", value=str(defaults['Số Hợp đồng thiết kế']), key=f"sohdtk_{wid_key}")
            ngay_hdtk = st.date_input("Ngày HĐ TK", value=defaults['Ngày Hợp đồng thiết kế'], format="DD/MM/YYYY", key=f"ngayhdtk_{wid_key}")
            gt_hdtk_str = st.text_input("Giá trị HĐ TK", value=format_num_val(defaults['Giá trị Hợp đồng thiết kế']), key=f"gthdtk_{wid_key}")
            gt_hdtk = parse_num_val(gt_hdtk_str)
        with c2:
            st.markdown("**HĐ Giám sát**")
            so_hdgs = st.text_input("Số HĐ GS", value=str(defaults['Số Hợp đồng giám sát']), key=f"sohdgs_{wid_key}")
            ngay_hdgs = st.date_input("Ngày HĐ GS", value=defaults['Ngày Hợp đồng giám sát'], format="DD/MM/YYYY", key=f"ngayhdgs_{wid_key}")
            gt_hdgs_str = st.text_input("Giá trị HĐ GS", value=format_num_val(defaults['Giá trị Hợp đồng giám sát']), key=f"gthdgs_{wid_key}")
            gt_hdgs = parse_num_val(gt_hdgs_str)
        with c3:
            st.markdown("**HĐ Xây lắp**")
            so_hdxl = st.text_input("Số HĐ XL", value=str(defaults['Số Hợp đồng xây lắp']), key=f"sohdxl_{wid_key}")
            ngay_hdxl = st.date_input("Ngày HĐ XL", value=defaults['Ngày Hợp đồng xây lắp'], format="DD/MM/YYYY", key=f"ngayhdxl_{wid_key}")
            gt_hdxl_str = st.text_input("Giá trị HĐ XL", value=format_num_val(defaults['Giá trị Hợp đồng xây lắp']), key=f"gthdxl_{wid_key}")
            gt_hdxl = parse_num_val(gt_hdxl_str)
        
        c4, c5, c6 = st.columns(3)
        with c4:
            gt_ps_str = st.text_input("Giá trị phát sinh", value=format_num_val(defaults['Giá trị phát sinh']), key=f"gtps_{wid_key}")
            gt_ps = parse_num_val(gt_ps_str)
        with c5:
            gt_vtt_str = st.text_input("Giá trị VT thừa", value=format_num_val(defaults['Giá trị VT thừa']), key=f"gtvtt_{wid_key}")
            gt_vtt = parse_num_val(gt_vtt_str)
        with c6:
            gt_vtth_str = st.text_input("Giá trị VTTH", value=format_num_val(defaults['Giá trị VTTH']), key=f"gtvtth_{wid_key}")
            gt_vtth = parse_num_val(gt_vtth_str)

    st.header("2. BẢNG TỔNG HỢP QUYẾT TOÁN KINH PHÍ SỬA CHỮA LỚN")
    st.write("Thêm các dòng cho hạng mục con. Cột 'Tên Hạng mục' sẽ được tự động map vào 'Tên Công trình' trong Excel theo cấu trúc mẫu.")
    st.info("LƯU Ý: Các mục tổng (A, B, C, E, SCL) và các mục con phụ thuộc sẽ tự động cộng dựa trên dữ liệu các hạng mục chi tiết (Ví dụ: A.1.1, B.1.1). Việc tính toán tuân theo quy tắc bạn đã đề ra.")
    
    if 'sub_df' not in st.session_state:
        initial_data = [
            {"STT": "A", "Tên Hạng mục": "CHI PHÍ VẬT TƯ, THIẾT BỊ (sau thuế)", "Giá trị Dự toán": 0, "Giá trị quyết toán": 0, "Chênh lệch": 0},
            {"STT": "A.1", "Tên Hạng mục": "Chi phí thiết bị", "Giá trị Dự toán": 0, "Giá trị quyết toán": 0, "Chênh lệch": 0},
            {"STT": "A.1.1", "Tên Hạng mục": "Thiết bị nhập khẩu", "Giá trị Dự toán": 0, "Giá trị quyết toán": 0, "Chênh lệch": 0},
            {"STT": "A.1.2", "Tên Hạng mục": "VT A cấp", "Giá trị Dự toán": 0, "Giá trị quyết toán": 0, "Chênh lệch": 0},
            {"STT": "A.1.3", "Tên Hạng mục": "Chi phí tháo dỡ, lắp đặt", "Giá trị Dự toán": 0, "Giá trị quyết toán": 0, "Chênh lệch": 0},
            {"STT": "A.1.4", "Tên Hạng mục": "Chi phí thí nghiệm, hiệu chỉnh", "Giá trị Dự toán": 0, "Giá trị quyết toán": 0, "Chênh lệch": 0},
            {"STT": "A.2", "Tên Hạng mục": "Chi phí vật tư", "Giá trị Dự toán": 0, "Giá trị quyết toán": 0, "Chênh lệch": 0},
            {"STT": "A.3", "Tên Hạng mục": "Thuế GTGT", "Giá trị Dự toán": 0, "Giá trị quyết toán": 0, "Chênh lệch": 0},
            {"STT": "B", "Tên Hạng mục": "CHI PHÍ SỬA CHỮA", "Giá trị Dự toán": 0, "Giá trị quyết toán": 0, "Chênh lệch": 0},
            {"STT": "B.1", "Tên Hạng mục": "Chi phí vật liệu", "Giá trị Dự toán": 0, "Giá trị quyết toán": 0, "Chênh lệch": 0},
            {"STT": "B.1.1", "Tên Hạng mục": "Vật liệu phần không áp dụng đơn giá XDCB", "Giá trị Dự toán": 0, "Giá trị quyết toán": 0, "Chênh lệch": 0},
            {"STT": "B.1.2", "Tên Hạng mục": "Vật liệu phần áp dụng đơn giá XDCB", "Giá trị Dự toán": 0, "Giá trị quyết toán": 0, "Chênh lệch": 0},
            {"STT": "B.1.3", "Tên Hạng mục": "Chênh lệch giá vật liệu phần áp dụng đơn giá XDCB", "Giá trị Dự toán": 0, "Giá trị quyết toán": 0, "Chênh lệch": 0},
            {"STT": "B.1.4", "Tên Hạng mục": "Vật liệu phụ trong SCL thiết bị", "Giá trị Dự toán": 0, "Giá trị quyết toán": 0, "Chênh lệch": 0},
            {"STT": "B.2", "Tên Hạng mục": "Chi phí nhân công", "Giá trị Dự toán": 0, "Giá trị quyết toán": 0, "Chênh lệch": 0},
            {"STT": "B.2.1", "Tên Hạng mục": "Chi phí nhân công phần không áp dụng đơn giá XDCB", "Giá trị Dự toán": 0, "Giá trị quyết toán": 0, "Chênh lệch": 0},
            {"STT": "B.2.2", "Tên Hạng mục": "Chi phí nhân công phần áp dụng đơn giá XDCB", "Giá trị Dự toán": 0, "Giá trị quyết toán": 0, "Chênh lệch": 0},
            {"STT": "B.3", "Tên Hạng mục": "Chi phí máy thi công", "Giá trị Dự toán": 0, "Giá trị quyết toán": 0, "Chênh lệch": 0},
            {"STT": "B.3.1", "Tên Hạng mục": "Chi phí máy thi công phần không áp dụng đơn giá XDCB", "Giá trị Dự toán": 0, "Giá trị quyết toán": 0, "Chênh lệch": 0},
            {"STT": "B.3.2", "Tên Hạng mục": "Chi phí máy thi công phần áp dụng đơn giá XDCB", "Giá trị Dự toán": 0, "Giá trị quyết toán": 0, "Chênh lệch": 0},
            {"STT": "B.4", "Tên Hạng mục": "Chi phí làm đêm, làm thêm giờ", "Giá trị Dự toán": 0, "Giá trị quyết toán": 0, "Chênh lệch": 0},
            {"STT": "B.5", "Tên Hạng mục": "Chi phí chung", "Giá trị Dự toán": 0, "Giá trị quyết toán": 0, "Chênh lệch": 0},
            {"STT": "B.6", "Tên Hạng mục": "Thu nhập chịu thuế tính trước", "Giá trị Dự toán": 0, "Giá trị quyết toán": 0, "Chênh lệch": 0},
            {"STT": "B.7", "Tên Hạng mục": "Giá trị sửa chữa trước thuế", "Giá trị Dự toán": 0, "Giá trị quyết toán": 0, "Chênh lệch": 0},
            {"STT": "B.8", "Tên Hạng mục": "Thuế GTGT", "Giá trị Dự toán": 0, "Giá trị quyết toán": 0, "Chênh lệch": 0},
            {"STT": "C", "Tên Hạng mục": "CHI PHÍ KHÁC (sau thuế)", "Giá trị Dự toán": 0, "Giá trị quyết toán": 0, "Chênh lệch": 0},
            {"STT": "C.1", "Tên Hạng mục": "Chi phí giám sát thi công xây dựng", "Giá trị Dự toán": 0, "Giá trị quyết toán": 0, "Chênh lệch": 0},
            {"STT": "C.2", "Tên Hạng mục": "Chi phí giám sát lắp đặt thiết bị", "Giá trị Dự toán": 0, "Giá trị quyết toán": 0, "Chênh lệch": 0},
            {"STT": "C.3", "Tên Hạng mục": "Chi phí bảo hiểm công trình", "Giá trị Dự toán": 0, "Giá trị quyết toán": 0, "Chênh lệch": 0},
            {"STT": "C.4", "Tên Hạng mục": "Chi phí thẩm tra - phê duyệt quyết toán", "Giá trị Dự toán": 0, "Giá trị quyết toán": 0, "Chênh lệch": 0},
            {"STT": "C.5", "Tên Hạng mục": "Vận chuyển VTTB A cấp đến công trường", "Giá trị Dự toán": 0, "Giá trị quyết toán": 0, "Chênh lệch": 0},
            {"STT": "C.6", "Tên Hạng mục": "Thuế GTGT", "Giá trị Dự toán": 0, "Giá trị quyết toán": 0, "Chênh lệch": 0},
            {"STT": "D", "Tên Hạng mục": "CHI PHÍ DỰ PHÒNG", "Giá trị Dự toán": 0, "Giá trị quyết toán": 0, "Chênh lệch": 0},
            {"STT": "E", "Tên Hạng mục": "Tổng giá trị sau thuế", "Giá trị Dự toán": 0, "Giá trị quyết toán": 0, "Chênh lệch": 0},
            {"STT": "E.1", "Tên Hạng mục": "Tổng giá trị trước thuế", "Giá trị Dự toán": 0, "Giá trị quyết toán": 0, "Chênh lệch": 0},
            {"STT": "E.2", "Tên Hạng mục": "Thuế GTGT", "Giá trị Dự toán": 0, "Giá trị quyết toán": 0, "Chênh lệch": 0},
            {"STT": "F", "Tên Hạng mục": "GIÁ TRỊ VẬT TƯ THU HỒI", "Giá trị Dự toán": 0, "Giá trị quyết toán": 0, "Chênh lệch": 0},
            {"STT": "SCL", "Tên Hạng mục": "CHI PHÍ SCL", "Giá trị Dự toán": 0, "Giá trị quyết toán": 0, "Chênh lệch": 0}
        ]
        st.session_state.sub_df = pd.DataFrame(initial_data, columns=sub_items_columns)
    
    col_config = {
        "Giá trị Dự toán": st.column_config.NumberColumn(format="%,d", step=1),
        "Giá trị quyết toán": st.column_config.NumberColumn(format="%,d", step=1),
        "Chênh lệch": st.column_config.NumberColumn(format="%,d", disabled=True)
    }
    edited_sub_df = st.data_editor(st.session_state.sub_df, num_rows="dynamic", use_container_width=True, column_config=col_config)

    # --- TÍNH TOÁN TỰ ĐỘNG ---
    calculated_df = edited_sub_df.copy()
    
    # Ép kiểu an toàn để tính toán
    for col in ['Giá trị Dự toán', 'Giá trị quyết toán']:
        calculated_df[col] = pd.to_numeric(calculated_df[col], errors='coerce').fillna(0).astype(int)

    calculated_df['Chênh lệch'] = calculated_df['Giá trị Dự toán'] - calculated_df['Giá trị quyết toán']

    changed = False

    def set_val(target_stt, val_dt, val_qt):
        global changed
        mask = calculated_df['STT'].astype(str).str.strip() == target_stt
        if mask.any():
            idx = mask.idxmax()
            if int(calculated_df.at[idx, 'Giá trị Dự toán']) != int(val_dt):
                calculated_df.at[idx, 'Giá trị Dự toán'] = int(val_dt)
                changed = True
            if int(calculated_df.at[idx, 'Giá trị quyết toán']) != int(val_qt):
                calculated_df.at[idx, 'Giá trị quyết toán'] = int(val_qt)
                changed = True

    def get_val(target_stt):
        mask = calculated_df['STT'].astype(str).str.strip() == target_stt
        if mask.any():
            idx = mask.idxmax()
            return int(calculated_df.at[idx, 'Giá trị Dự toán']), int(calculated_df.at[idx, 'Giá trị quyết toán'])
        return 0, 0

    def get_sum(stt_list):
        s_dt, s_qt = 0, 0
        for stt in stt_list:
            dt, qt = get_val(stt)
            s_dt += dt
            s_qt += qt
        return s_dt, s_qt

    # A: 1=1.1+1.2+1.3+1.4
    val_A1_dt, val_A1_qt = get_sum(['A.1.1', 'A.1.2', 'A.1.3', 'A.1.4'])
    set_val('A.1', val_A1_dt, val_A1_qt)
    
    val_A2_dt, val_A2_qt = get_val('A.2')
    
    # A: 3=(1+2)*10%
    val_A3_dt = int(round((val_A1_dt + val_A2_dt) * 0.10))
    val_A3_qt = int(round((val_A1_qt + val_A2_qt) * 0.10))
    set_val('A.3', val_A3_dt, val_A3_qt)
    
    # A = 1+2+3
    val_A_dt = val_A1_dt + val_A2_dt + val_A3_dt
    val_A_qt = val_A1_qt + val_A2_qt + val_A3_qt
    set_val('A', val_A_dt, val_A_qt)

    # B: 1=1.1+1.2+1.3+1.4, 2=2.1+2.2, 3=3.1+3.2
    val_B1_dt, val_B1_qt = get_sum(['B.1.1', 'B.1.2', 'B.1.3', 'B.1.4'])
    set_val('B.1', val_B1_dt, val_B1_qt)
    
    val_B2_dt, val_B2_qt = get_sum(['B.2.1', 'B.2.2'])
    set_val('B.2', val_B2_dt, val_B2_qt)
    
    val_B3_dt, val_B3_qt = get_sum(['B.3.1', 'B.3.2'])
    set_val('B.3', val_B3_dt, val_B3_qt)
    
    val_B4_dt, val_B4_qt = get_val('B.4')
    val_B5_dt, val_B5_qt = get_val('B.5')
    val_B6_dt, val_B6_qt = get_val('B.6')

    # B.7 = 1+2+3+4+5+6
    val_B7_dt = sum([val_B1_dt, val_B2_dt, val_B3_dt, val_B4_dt, val_B5_dt, val_B6_dt])
    val_B7_qt = sum([val_B1_qt, val_B2_qt, val_B3_qt, val_B4_qt, val_B5_qt, val_B6_qt])
    set_val('B.7', val_B7_dt, val_B7_qt)
    
    # B.8 = 7*8%
    val_B8_dt = int(round(val_B7_dt * 0.08))
    val_B8_qt = int(round(val_B7_qt * 0.08))
    set_val('B.8', val_B8_dt, val_B8_qt)
    
    # B = 7+8
    val_B_dt = val_B7_dt + val_B8_dt
    val_B_qt = val_B7_qt + val_B8_qt
    set_val('B', val_B_dt, val_B_qt)

    # C: 6=(1+2+3+4+5)*8%
    val_C1_dt, val_C1_qt = get_val('C.1')
    val_C2_dt, val_C2_qt = get_val('C.2')
    val_C3_dt, val_C3_qt = get_val('C.3')
    val_C4_dt, val_C4_qt = get_val('C.4')
    val_C5_dt, val_C5_qt = get_val('C.5')
    sum_C_1_5_dt = val_C1_dt + val_C2_dt + val_C3_dt + val_C4_dt + val_C5_dt
    sum_C_1_5_qt = val_C1_qt + val_C2_qt + val_C3_qt + val_C4_qt + val_C5_qt
    
    val_C6_dt = int(round(sum_C_1_5_dt * 0.08))
    val_C6_qt = int(round(sum_C_1_5_qt * 0.08))
    set_val('C.6', val_C6_dt, val_C6_qt)
    
    # C = 1+2+3+4+5+6
    val_C_dt = sum_C_1_5_dt + val_C6_dt
    val_C_qt = sum_C_1_5_qt + val_C6_qt
    set_val('C', val_C_dt, val_C_qt)

    # D là ô nhập liệu
    val_D_dt, val_D_qt = get_val('D')

    # E.1 = 1+2 (A) + 7 (B) + (1+2+3+4+5) (C) + D
    val_E1_dt = val_A1_dt + val_A2_dt + val_B7_dt + sum_C_1_5_dt + val_D_dt
    val_E1_qt = val_A1_qt + val_A2_qt + val_B7_qt + sum_C_1_5_qt + val_D_qt
    set_val('E.1', val_E1_dt, val_E1_qt)
    
    # E.2 = 3 (A) + 8 (B) + 6 (C)
    val_E2_dt = val_A3_dt + val_B8_dt + val_C6_dt
    val_E2_qt = val_A3_qt + val_B8_qt + val_C6_qt
    set_val('E.2', val_E2_dt, val_E2_qt)
    
    # E = E.1 + E.2
    val_E_dt = val_E1_dt + val_E2_dt
    val_E_qt = val_E1_qt + val_E2_qt
    set_val('E', val_E_dt, val_E_qt)

    val_F_dt, val_F_qt = get_val('F')
    
    # MỤC CUỐI CÙNG = E.1 - F
    val_SCL_dt = val_E1_dt - val_F_dt
    val_SCL_qt = val_E1_qt - val_F_qt
    set_val('SCL', val_SCL_dt, val_SCL_qt)


    # 4. Lưu lại và yêu cầu nạp lại ứng dụng nếu có thay đổi tự động
    if changed:
        st.session_state.sub_df = calculated_df
        try:
            st.rerun()
        except AttributeError:
            st.experimental_rerun()


    if st.button("💾 Lưu trữ dữ liệu", type="primary"):
        if not ten_ct.strip():
            st.error("Vui lòng nhập Tên Công trình (bắt buộc)!")
        else:
            try:
                # 1. Row Công trình chính
                main_row = {
                    'STT': stt, 'Tên Công trình': ten_ct, 'Mã CT': ma_ct, 'Kế hoạch': ke_hoach,
                    'Số Phương án': so_pa, 'Ngày Phương án': ngay_pa, 'Giá trị Phương án': gt_pa,
                    'Số Dự toán': so_dt, 'Ngày Dự toán': ngay_dt, 'Giá trị Dự toán': gt_dt,
                    'Số Hợp đồng thiết kế': so_hdtk, 'Ngày Hợp đồng thiết kế': ngay_hdtk, 'Giá trị Hợp đồng thiết kế': gt_hdtk,
                    'Số Hợp đồng giám sát': so_hdgs, 'Ngày Hợp đồng giám sát': ngay_hdgs, 'Giá trị Hợp đồng giám sát': gt_hdgs,
                    'Số Hợp đồng xây lắp': so_hdxl, 'Ngày Hợp đồng xây lắp': ngay_hdxl, 'Giá trị Hợp đồng xây lắp': gt_hdxl,
                    'Giá trị phát sinh': gt_ps, 'Giá trị VT thừa': gt_vtt, 'Giá trị VTTH': gt_vtth,
                    'Số Q.định phê duyệt QT công trình': so_qd, 'Ngày Q.định phê duyệt QT công trình': ngay_qd, 
                    'Giá trị Q.định phê duyệt QT công trình': gt_qd, 'Số tiền bằng chữ': so_tien_chu, 
                    'Ghi chú': ghi_chu, 'Đơn vị QL': don_vi_ql,
                    'Căn cứ pháp lý': can_cu_phap_ly, 'Khối lượng công việc': khoi_luong_cv,
                    'Ngày khởi công': ngay_khoi_cong, 'Ngày hoàn thành': ngay_hoan_thanh
                }
                
                rows_to_add = [main_row]
                
                # 2. Rows Hạng mục con
                for index, row in edited_sub_df.iterrows():
                    if pd.notna(row['Tên Hạng mục']) and str(row['Tên Hạng mục']).strip() != "":
                        sub_row = {col: None for col in ALL_COLUMNS} # Fill None trước
                        sub_row['STT'] = row.get('STT')
                        sub_row['Tên Công trình'] = row.get('Tên Hạng mục') # Hạng mục con ghi vào chung cột
                        sub_row['Giá trị Dự toán'] = row.get('Giá trị Dự toán')
                        sub_row['Giá trị Q.định phê duyệt QT công trình'] = row.get('Giá trị quyết toán')
                        sub_row['Ghi chú'] = row.get('Chênh lệch')
                        # Gắn thêm một cờ/mã để lọc theo Công trình cha nếu cần (tuỳ chọn)
                        # sub_row['Mã CT'] = ma_ct # Cách linh hoạt để báo cáo biết nó thuộc công trình nào
                        rows_to_add.append(sub_row)
                
                new_data = pd.DataFrame(rows_to_add)
                
                # Mở file hiện tại (hoặc tạo mới nếu mất)
                db_df = load_data()
                
                if selected_edit_ct != "-- Thêm mới --":
                    start_indices = db_df.index[db_df['Tên Công trình'] == selected_edit_ct].tolist()
                    if start_indices:
                        start_idx = start_indices[0]
                        end_idx = len(db_df)
                        for i in range(start_idx + 1, len(db_df)):
                            val = str(db_df.at[i, 'STT']).strip().upper()
                            if val in ['I', 'II', 'III', 'IV', 'V', 'VI', 'VII', 'VIII', 'IX', 'X']:
                                end_idx = i
                                break
                        db_df = db_df.drop(db_df.index[start_idx:end_idx])
                
                # Nối dữ liệu
                updated_df = pd.concat([db_df, new_data], ignore_index=True)
                
                # Lưu đè lại
                updated_df.to_excel(DB_FILE, index=False)
                
                if 'sub_df' in st.session_state:
                    del st.session_state['sub_df']
                st.session_state.last_edit_ct = "-- Thêm mới --"
                
                st.success("🎉 Lưu trữ dữ liệu thành công!")
                
            except Exception as e:
                st.error(f"Lỗi khi lưu dữ liệu: {e}")

with tab2:
    st.header("📊 Báo cáo Công trình")
    db_df = load_data()
    
    if db_df.empty:
        st.info("Chưa có dữ liệu. Vui lòng nhập liệu ở tab 'Nhập liệu'.")
    else:
        # Lấy danh sách các Công trình chính (các row có Kế hoạch khác NaN/None)
        main_mask = db_df['Kế hoạch'].notna()
        list_cong_trinh = db_df.loc[main_mask, 'Tên Công trình'].dropna().unique().tolist()
        
        selected_ct = st.selectbox("Chọn Công trình để xem báo cáo:", ["-- Chọn --", "Tất cả các công trình"] + list_cong_trinh)
        
        if selected_ct != "-- Chọn --":
            if selected_ct == "Tất cả các công trình":
                st.dataframe(db_df, use_container_width=True)
                
                output = BytesIO()
                import openpyxl
                import copy
                
                try:
                    wb = openpyxl.load_workbook('Mẫu 04.xlsx')
                except Exception as e:
                    st.error(f"Không thể mở file Mẫu 04.xlsx. Vui lòng kiểm tra lại file. Lỗi: {e}")
                    st.stop()
                    
                template_sheet = wb.active
                template_sheet.title = 'Template_04'
                
                row_map = {
                    "A": 11, "A.1": 12, "A.1.1": 13, "A.1.2": 14, "A.1.3": 15, "A.1.4": 16, "A.2": 17, "A.3": 18,
                    "B": 19, "B.1": 20, "B.1.1": 21, "B.1.2": 22, "B.1.3": 23, "B.1.4": 24, "B.2": 25, "B.2.1": 26, "B.2.2": 27,
                    "B.3": 28, "B.3.1": 29, "B.3.2": 30, "B.4": 31, "B.5": 32, "B.6": 33, "B.7": 34, "B.8": 35,
                    "C": 36, "C.1": 37, "C.2": 38, "C.3": 39, "C.4": 40, "C.5": 41, "C.6": 42,
                    "D": 46, "E": 47, "E.1": 48, "E.2": 49, "F": 50, "SCL": 51
                }
                
                for ct in list_cong_trinh:
                    start_indices = db_df.index[db_df['Tên Công trình'] == ct].tolist()
                    if start_indices:
                        start_idx = start_indices[0]
                        end_idx = len(db_df)
                        for i in range(start_idx + 1, len(db_df)):
                            val = str(db_df.at[i, 'STT']).strip().upper()
                            if val in ['I', 'II', 'III', 'IV', 'V', 'VI', 'VII', 'VIII', 'IX', 'X']:
                                end_idx = i
                                break
                        
                        report_df = db_df.iloc[start_idx:end_idx]
                        safe_sheet_name = str(ct)[:31].replace(":", "").replace("/", "").replace("\\", "").replace("?", "").replace("*", "").replace("[", "").replace("]", "")
                        if not safe_sheet_name: safe_sheet_name = "Sheet"
                        
                        ws = wb.copy_worksheet(template_sheet)
                        ws.title = safe_sheet_name
                        ws['D8'] = ct # Luôn lấy Tên công trình
                        
                        # CẬP NHẬT TÊN MỤC E, E.1, E.2 CHO KHỚP VỚI CÁCH TÍNH MỚI
                        ws.cell(row=47, column=2).value = "Tổng giá trị sau thuế"
                        ws.cell(row=47, column=3).value = "Tổng giá trị sau thuế"
                        ws.cell(row=48, column=2).value = "Tổng giá trị trước thuế"
                        ws.cell(row=48, column=3).value = "Tổng giá trị trước thuế"
                        ws.cell(row=49, column=2).value = "Thuế GTGT"
                        ws.cell(row=49, column=3).value = "Thuế GTGT"
                        
                        for idx, row in report_df.iterrows():
                            stt_val = str(row['STT']).strip().upper()
                            if stt_val in row_map:
                                r_idx = row_map[stt_val]
                                dt = row['Giá trị Dự toán']
                                qt = row['Giá trị Q.định phê duyệt QT công trình']
                                if pd.notna(dt):
                                    ws.cell(row=r_idx, column=4).value = float(dt) if dt else 0 # DỰ TOÁN
                                if pd.notna(qt):
                                    ws.cell(row=r_idx, column=7).value = float(qt) if qt else 0 # GIÁ TRỊ QUYẾT TOÁN

                wb.remove(template_sheet)
                if 'Sheet' in wb.sheetnames and len(wb.sheetnames) > 1:
                    wb.remove(wb['Sheet'])
                    
                wb.save(output)
                processed_data = output.getvalue()
                st.download_button(label="Tải xuống báo cáo (Excel)", data=processed_data, file_name="Bao_cao_Tat_ca_Cong_trinh.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
                
            else:
                start_indices = db_df.index[db_df['Tên Công trình'] == selected_ct].tolist()
                if start_indices:
                    start_idx = start_indices[0]
                    end_idx = len(db_df)
                    for i in range(start_idx + 1, len(db_df)):
                        val = str(db_df.at[i, 'STT']).strip().upper()
                        if val in ['I', 'II', 'III', 'IV', 'V', 'VI', 'VII', 'VIII', 'IX', 'X']:
                            end_idx = i
                            break
                            
                    report_df = db_df.iloc[start_idx:end_idx]
                    st.dataframe(report_df, use_container_width=True)
                    
                    output = BytesIO()
                    import openpyxl
                    try:
                        wb = openpyxl.load_workbook('Mẫu 04.xlsx')
                    except Exception as e:
                        st.error(f"Không thể mở file Mẫu 04.xlsx. Lỗi: {e}")
                        st.stop()
                        
                    ws = wb.active
                    ws.title = 'Báo cáo'
                    ws['D8'] = selected_ct
                    
                    # CẬP NHẬT TÊN MỤC E, E.1, E.2 CHO KHỚP VỚI CÁCH TÍNH MỚI
                    ws.cell(row=47, column=2).value = "Tổng giá trị sau thuế"
                    ws.cell(row=47, column=3).value = "Tổng giá trị sau thuế"
                    ws.cell(row=48, column=2).value = "Tổng giá trị trước thuế"
                    ws.cell(row=48, column=3).value = "Tổng giá trị trước thuế"
                    ws.cell(row=49, column=2).value = "Thuế GTGT"
                    ws.cell(row=49, column=3).value = "Thuế GTGT"
                    
                    row_map = {
                        "A": 11, "A.1": 12, "A.1.1": 13, "A.1.2": 14, "A.1.3": 15, "A.1.4": 16, "A.2": 17, "A.3": 18,
                        "B": 19, "B.1": 20, "B.1.1": 21, "B.1.2": 22, "B.1.3": 23, "B.1.4": 24, "B.2": 25, "B.2.1": 26, "B.2.2": 27,
                        "B.3": 28, "B.3.1": 29, "B.3.2": 30, "B.4": 31, "B.5": 32, "B.6": 33, "B.7": 34, "B.8": 35,
                        "C": 36, "C.1": 37, "C.2": 38, "C.3": 39, "C.4": 40, "C.5": 41, "C.6": 42,
                        "D": 46, "E": 47, "E.1": 48, "E.2": 49, "F": 50, "SCL": 51
                    }
                    
                    for idx, row in report_df.iterrows():
                        stt_val = str(row['STT']).strip().upper()
                        if stt_val in row_map:
                            r_idx = row_map[stt_val]
                            dt = row['Giá trị Dự toán']
                            qt = row['Giá trị Q.định phê duyệt QT công trình']
                            if pd.notna(dt):
                                ws.cell(row=r_idx, column=4).value = float(dt) if dt else 0
                            if pd.notna(qt):
                                ws.cell(row=r_idx, column=7).value = float(qt) if qt else 0
                    
                    wb.save(output)
                    processed_data = output.getvalue()
                    st.download_button(label="Tải xuống báo cáo (Excel)", data=processed_data, file_name=f"Bao_cao_{selected_ct[:20]}.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

with tab3:
    st.header("📄 Bảng thuyết minh quyết toán")
    db_df_tab3 = load_data()
    
    if db_df_tab3.empty:
        st.info("Chưa có dữ liệu. Vui lòng nhập liệu ở tab 'Nhập liệu'.")
    else:
        main_mask_tab3 = db_df_tab3['Kế hoạch'].notna()
        list_ct_tab3 = db_df_tab3.loc[main_mask_tab3, 'Tên Công trình'].dropna().unique().tolist()
        
        selected_ct_tab3 = st.selectbox("Chọn Công trình để xuất Thuyết minh QT:", ["-- Chọn --"] + list_ct_tab3, key="tmqt_select")
        
        if selected_ct_tab3 != "-- Chọn --":
            # Lấy dữ liệu công trình
            start_indices_t3 = db_df_tab3.index[db_df_tab3['Tên Công trình'] == selected_ct_tab3].tolist()
            if start_indices_t3:
                start_idx_t3 = start_indices_t3[0]
                end_idx_t3 = len(db_df_tab3)
                for i in range(start_idx_t3 + 1, len(db_df_tab3)):
                    val = str(db_df_tab3.at[i, 'STT']).strip().upper()
                    if val in ['I', 'II', 'III', 'IV', 'V', 'VI', 'VII', 'VIII', 'IX', 'X']:
                        end_idx_t3 = i
                        break
                
                ct_data_t3 = db_df_tab3.iloc[start_idx_t3:end_idx_t3]
                main_row_t3 = ct_data_t3.iloc[0]
                
                # Lấy các giá trị cần thiết
                ten_ct_t3 = str(main_row_t3.get('Tên Công trình', ''))
                ma_ct_t3 = str(main_row_t3.get('Mã CT', '')) if pd.notna(main_row_t3.get('Mã CT')) else ''
                ke_hoach_t3 = main_row_t3.get('Kế hoạch', 0)
                if pd.isna(ke_hoach_t3): ke_hoach_t3 = 0
                ke_hoach_t3 = int(float(ke_hoach_t3))
                
                don_vi_ql_t3 = str(main_row_t3.get('Đơn vị QL', '')) if pd.notna(main_row_t3.get('Đơn vị QL')) else ''
                
                # Lấy căn cứ pháp lý và khối lượng công việc
                can_cu_pl_t3 = str(main_row_t3.get('Căn cứ pháp lý', '')) if pd.notna(main_row_t3.get('Căn cứ pháp lý')) else ''
                klcv_t3 = str(main_row_t3.get('Khối lượng công việc', '')) if pd.notna(main_row_t3.get('Khối lượng công việc')) else ''
                
                # Lấy ngày khởi công và hoàn thành
                ngay_kc_t3 = main_row_t3.get('Ngày khởi công')
                ngay_ht_t3 = main_row_t3.get('Ngày hoàn thành')
                
                def format_date_vn(d):
                    if pd.isna(d) or d is None:
                        return '....../....../...........'
                    if isinstance(d, pd.Timestamp):
                        d = d.date()
                    if isinstance(d, (datetime.date, datetime.datetime)):
                        return d.strftime('%d/%m/%Y')
                    return str(d)
                
                ngay_kc_str = format_date_vn(ngay_kc_t3)
                ngay_ht_str = format_date_vn(ngay_ht_t3)
                
                # Lấy giá trị SCL dự toán và quyết toán từ bảng tổng hợp
                gt_dt_scl = 0
                gt_qt_scl = 0
                for idx_t3, row_t3 in ct_data_t3.iterrows():
                    stt_val = str(row_t3['STT']).strip().upper()
                    if stt_val == 'SCL':
                        dt_val = row_t3.get('Giá trị Dự toán', 0)
                        qt_val = row_t3.get('Giá trị Q.định phê duyệt QT công trình', 0)
                        if pd.notna(dt_val): gt_dt_scl = int(float(dt_val))
                        if pd.notna(qt_val): gt_qt_scl = int(float(qt_val))
                        break
                
                # Hiển thị preview
                st.subheader("Xem trước thông tin")
                preview_col1, preview_col2 = st.columns(2)
                with preview_col1:
                    st.write(f"**Tên công trình:** {ten_ct_t3}")
                    st.write(f"**MCT:** {ma_ct_t3}")
                    st.write(f"**Giá trị kế hoạch vốn:** {f'{ke_hoach_t3:,}'} đồng")
                    st.write(f"**Giá trị dự toán được duyệt:** {f'{gt_dt_scl:,}'} đồng")
                with preview_col2:
                    st.write(f"**Thời gian khởi công:** {ngay_kc_str}")
                    st.write(f"**Thời gian hoàn thành:** {ngay_ht_str}")
                    st.write(f"**Giá trị quyết toán:** {f'{gt_qt_scl:,}'} đồng")
                    st.write(f"**Đơn vị QL:** {don_vi_ql_t3}")
                
                if can_cu_pl_t3:
                    st.write("**Căn cứ pháp lý:**")
                    st.text(can_cu_pl_t3)
                if klcv_t3:
                    st.write("**Khối lượng công việc:**")
                    st.text(klcv_t3)
                
                st.divider()
                
                if st.button("📥 Xuất file Word - Bảng thuyết minh quyết toán", type="primary", key="btn_tmqt"):
                    try:
                        doc = DocxDocument('Mẫu TMQT.docx')
                        
                        ghi_chu_t3 = str(main_row_t3.get('Ghi chú', '')) if pd.notna(main_row_t3.get('Ghi chú')) else ''
                        now = datetime.datetime.now()
                        
                        def format_money(val):
                            if val == 0: return '0'
                            return f'{val:,}'.replace(',', '.')
                        
                        def format_trieu(val):
                            if val == 0: return '0'
                            trieu = val / 1000000
                            if trieu == int(trieu):
                                return f'{int(trieu):,}'.replace(',', '.')
                            return f'{trieu:,.2f}'.replace(',', '.')
                        
                        # Set paragraph spacing to 0 and replace text
                        for p in doc.paragraphs:
                            p.paragraph_format.space_after = Pt(0)
                            
                            text_val = p.text.strip()
                            if not text_val: continue
                            
                            if "- Tên danh mục:" in text_val:
                                p.text = f"- Tên danh mục: {ten_ct_t3}\n- Mã công trình: {ma_ct_t3}"
                            elif "- Giá trị vốn kế hoạch:" in text_val:
                                p.text = f"- Giá trị vốn kế hoạch: {format_money(ke_hoach_t3)} đồng"
                            elif "sửa chữa lớn năm" in text_val:
                                p.text = f"- Thuộc kế hoạch vốn sửa chữa lớn năm {now.year}"
                            elif "Hình thức tự làm hay thuê ngoài" in text_val:
                                p.text = f"- Hình thức tự làm hay thuê ngoài: {ghi_chu_t3}"
                            elif "- Tên đơn vị thi công" in text_val:
                                p.text = f"- Tên đơn vị thi công: {don_vi_ql_t3}"
                            elif "- Giá trị dự toán được duyệt" in text_val:
                                p.text = f"- Giá trị dự toán được duyệt: {format_money(gt_dt_scl)} đồng"
                            elif "- Thời gian khởi công" in text_val:
                                p.text = f"- Thời gian khởi công: {ngay_kc_str}"
                            elif "- Thời gian hoàn thành" in text_val:
                                p.text = f"- Thời gian hoàn thành: {ngay_ht_str}"
                            elif "- Giá trị quyết toán" in text_val and "hoàn thành" in text_val:
                                p.text = f"- Giá trị quyết toán danh mục hoàn thành: {format_money(gt_qt_scl)} đồng"
                            elif "Khối lượng công việc chủ yếu đã tiến hành" in text_val:
                                p.text = f"- Khối lượng công việc chủ yếu đã tiến hành (thay thế, sửa chữa những bộ phận nào của TSCĐ):"
                                if klcv_t3:
                                    p.text += f"\n{klcv_t3}"
                            elif "Các căn cứ về chế độ để lập quyết toán" in text_val:
                                p.text = f"- Các căn cứ về chế độ để lập quyết toán:"
                                if can_cu_pl_t3:
                                    p.text += f"\n{can_cu_pl_t3}"
                            elif "+ .........." in text_val:
                                p.text = ""
                            elif "ngày       tháng      năm" in text_val:
                                p.text = text_val.replace("2026", str(now.year))
                        
                        # Đảm bảo toàn bộ tài liệu là Times New Roman 12
                        if 'Normal' in doc.styles:
                            doc.styles['Normal'].font.name = 'Times New Roman'
                            doc.styles['Normal'].font.size = Pt(12)
                            
                        # Cài đặt Page Setup (Top, Bottom, Right = 2cm; Left = 3cm)
                        for section in doc.sections:
                            section.top_margin = Cm(2)
                            section.bottom_margin = Cm(2)
                            section.left_margin = Cm(3)
                            section.right_margin = Cm(2)
                            
                        for p in doc.paragraphs:
                            for run in p.runs:
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(12)
                                
                        for t in doc.tables:
                            for row in t.rows:
                                for cell in row.cells:
                                    for p in cell.paragraphs:
                                        p.paragraph_format.space_after = Pt(0)
                                        for run in p.runs:
                                            run.font.name = 'Times New Roman'
                                            run.font.size = Pt(12)
                        
                        # Lưu file Word vào BytesIO
                        output_docx = BytesIO()
                        doc.save(output_docx)
                        docx_data = output_docx.getvalue()
                        
                        safe_name = ten_ct_t3[:30].replace('/', '_').replace('\\', '_').replace(':', '_')
                        st.download_button(
                            label="📥 Tải xuống file Word",
                            data=docx_data,
                            file_name=f"Thuyet_minh_QT_{safe_name}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="download_tmqt"
                        )
                        st.success("✅ Đã tạo file Word thành công! Nhấn nút tải xuống bên trên.")
                        
                    except Exception as e:
                        st.error(f"Lỗi khi tạo file Word: {e}")
                        import traceback
                        st.code(traceback.format_exc())
