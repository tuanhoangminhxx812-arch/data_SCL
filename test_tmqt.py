import datetime
from docx import Document
from docx.shared import Pt

doc = Document("Mẫu TMQT.docx")
ten_ct_t3 = "TEST CONG TRINH"
ma_ct_t3 = "CT-001"
ke_hoach_t3 = 1500000000
ghi_chu_t3 = "Tu lam"
don_vi_ql_t3 = "AQL"
gt_dt_scl = 1200000000
ngay_kc_str = "01/01/2026"
ngay_ht_str = "01/12/2026"
gt_qt_scl = 1100000000
klcv_t3 = "- Cong viec 1\n- Cong viec 2"
can_cu_pl_t3 = "Quyet Dinh 1\nQuyet Dinh 2"

now = datetime.datetime.now()

def format_money(val):
    if val == 0: return '0'
    return f'{val:,}'.replace(',', '.')

def format_trieu(val):
    if val == 0: return '0'
    trieu = val / 1000000
    if trieu == int(trieu):
        return f'{int(trieu):,}'.replace(',', '.')
    return f'{trieu:,.2f}'.replace(',', '.')

for p in doc.paragraphs:
    p.paragraph_format.space_after = Pt(0)
    
    text_val = p.text.strip()
    if not text_val: continue
    
    if "- Tên danh mục:" in text_val:
        p.text = f"- Tên danh mục: {ten_ct_t3}\n- Mã công trình: {ma_ct_t3}"
    elif "- Giá trị vốn kế hoạch:" in text_val:
        p.text = f"- Giá trị vốn kế hoạch: {format_money(ke_hoach_t3)} đồng ({format_trieu(ke_hoach_t3)} triệu đồng)"
    elif "sửa chữa lớn năm" in text_val:
        p.text = f"- Thuộc kế hoạch vốn sửa chữa lớn năm {now.year}"
    elif "Hình thức tự làm hay thuê ngoài" in text_val:
        p.text = f"- Hình thức tự làm hay thuê ngoài: {ghi_chu_t3}"
    elif "- Tên đơn vị thi công" in text_val:
        p.text = f"- Tên đơn vị thi công: {don_vi_ql_t3}"
    elif "- Giá trị dự toán được duyệt" in text_val:
        p.text = f"- Giá trị dự toán được duyệt: {format_money(gt_dt_scl)} đồng"
    elif "- Thời gian khởi công" in text_val:
        p.text = f"- Thời gian khởi công: {ngay_kc_str}"
    elif "- Thời gian hoàn thành" in text_val:
        p.text = f"- Thời gian hoàn thành: {ngay_ht_str}"
    elif "- Giá trị quyết toán" in text_val and "hoàn thành" in text_val:
        p.text = f"- Giá trị quyết toán danh mục hoàn thành: {format_money(gt_qt_scl)} đồng"
    elif "Khối lượng công việc chủ yếu đã tiến hành" in text_val:
        p.text = f"- Khối lượng công việc chủ yếu đã tiến hành (thay thế, sửa chữa những bộ phận nào của TSCĐ):"
        if klcv_t3:
            p.text += f"\n{klcv_t3}"
    elif "Các căn cứ về chế độ để lập quyết toán" in text_val:
        p.text = f"- Các căn cứ về chế độ để lập quyết toán:"
        if can_cu_pl_t3:
            p.text += f"\n{can_cu_pl_t3}"
    elif "+ .........." in text_val:
        p.text = ""
    elif "ngày       tháng      năm" in text_val:
        p.text = text_val.replace("2026", str(now.year))

for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)

doc.save("test_tmqt_output.docx")
print("SUCCESS!")
