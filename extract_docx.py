import docx
import os

files = ["Mẫu TMQT.docx", "Mẫu QĐ QTCT.docx", "Mẫu phiếu báo cáo thẩm tra QT.docx", "Mẫu phiếu thẩm tra QT.docx"]
with open("extract_all_out.txt", "w", encoding="utf-8") as f:
    for file in files:
        if os.path.exists(file):
            f.write(f"\n\n--- FILE: {file} ---\n")
            doc = docx.Document(file)
            for p in doc.paragraphs:
                f.write(p.text + "\n")
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                             f.write(p.text + "\n")
