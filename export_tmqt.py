"""Script xuất file Word - Bảng thuyết minh quyết toán từ dữ liệu công trình đầu tiên"""
import pandas as pd
import datetime
from docx import Document as DocxDocument
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

DB_FILE = 'database_cong_trinh.xlsx'
db_df = pd.read_excel(DB_FILE)

# Lấy công trình đầu tiên
main_mask = db_df['Kế hoạch'].notna()
list_ct = db_df.loc[main_mask, 'Tên Công trình'].dropna().unique().tolist()
selected_ct = list_ct[0]
print(f"Xuất cho công trình: {selected_ct}")

start_idx = db_df.index[db_df['Tên Công trình'] == selected_ct].tolist()[0]
end_idx = len(db_df)
for i in range(start_idx + 1, len(db_df)):
    val = str(db_df.at[i, 'STT']).strip().upper()
    if val in ['I', 'II', 'III', 'IV', 'V', 'VI', 'VII', 'VIII', 'IX', 'X']:
        end_idx = i
        break

ct_data = db_df.iloc[start_idx:end_idx]
main_row = ct_data.iloc[0]

ten_ct = str(main_row.get('Tên Công trình', ''))
ma_ct = str(main_row.get('Mã CT', '')) if pd.notna(main_row.get('Mã CT')) else ''
ke_hoach = main_row.get('Kế hoạch', 0)
if pd.isna(ke_hoach): ke_hoach = 0
ke_hoach = int(float(ke_hoach))
don_vi_ql = str(main_row.get('Đơn vị QL', '')) if pd.notna(main_row.get('Đơn vị QL')) else ''
can_cu_pl = str(main_row.get('Căn cứ pháp lý', '')) if pd.notna(main_row.get('Căn cứ pháp lý')) else ''
klcv = str(main_row.get('Khối lượng công việc', '')) if pd.notna(main_row.get('Khối lượng công việc')) else ''

ngay_kc = main_row.get('Ngày khởi công')
ngay_ht = main_row.get('Ngày hoàn thành')

def format_date_vn(d):
    if pd.isna(d) or d is None:
        return '....../....../...........'
    if isinstance(d, pd.Timestamp):
        d = d.date()
    if isinstance(d, (datetime.date, datetime.datetime)):
        return d.strftime('%d/%m/%Y')
    return str(d)

ngay_kc_str = format_date_vn(ngay_kc)
ngay_ht_str = format_date_vn(ngay_ht)

gt_dt_scl = 0
gt_qt_scl = 0
for idx, row in ct_data.iterrows():
    stt_val = str(row['STT']).strip().upper()
    if stt_val == 'SCL':
        dt_val = row.get('Giá trị Dự toán', 0)
        qt_val = row.get('Giá trị Q.định phê duyệt QT công trình', 0)
        if pd.notna(dt_val): gt_dt_scl = int(float(dt_val))
        if pd.notna(qt_val): gt_qt_scl = int(float(qt_val))
        break

# === TẠO FILE WORD ===
doc = DocxDocument()

section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.5)
section.right_margin = Cm(1.5)

# === HEADER ===
header_table = doc.add_table(rows=2, cols=2)
header_table.alignment = WD_TABLE_ALIGNMENT.CENTER

cell_left = header_table.cell(0, 0)
p = cell_left.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("TỔNG CÔNG TY")
run.font.size = Pt(11)
run.font.name = 'Times New Roman'

cell_right = header_table.cell(0, 1)
p = cell_right.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
run.font.size = Pt(11)
run.font.name = 'Times New Roman'
run.font.bold = True

cell_left2 = header_table.cell(1, 0)
p = cell_left2.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("ĐIỆN LỰC TP HỒ CHÍ MINH TNHH")
run.font.size = Pt(11)
run.font.name = 'Times New Roman'
p2 = cell_left2.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("CÔNG TY ĐIỆN LỰC VŨNG TÀU")
run2.font.size = Pt(11)
run2.font.name = 'Times New Roman'
run2.font.bold = True

cell_right2 = header_table.cell(1, 1)
p = cell_right2.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Độc lập - Tự do - Hạnh phúc")
run.font.size = Pt(11)
run.font.name = 'Times New Roman'
run.font.bold = True
run.font.italic = True

# Xóa border bảng header
def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = tbl.makeelement(qn('w:tblPr'), {})
        tbl.insert(0, tblPr)
    borders = tblPr.find(qn('w:tblBorders'))
    if borders is not None:
        tblPr.remove(borders)
    borders = tblPr.makeelement(qn('w:tblBorders'), {})
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border_el = borders.makeelement(qn(f'w:{border_name}'), {
            qn('w:val'): 'none', qn('w:sz'): '0', qn('w:space'): '0', qn('w:color'): 'auto'
        })
        borders.append(border_el)
    tblPr.append(borders)

remove_table_borders(header_table)

# Ngày tháng
now = datetime.datetime.now()
p_date = doc.add_paragraph()
p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p_date.add_run(f"Vũng Tàu, ngày      tháng {now.month:02d}  năm {now.year}")
run.font.size = Pt(11)
run.font.name = 'Times New Roman'
run.font.italic = True

# === TIÊU ĐỀ ===
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.space_before = Pt(12)
p_title.space_after = Pt(12)
run = p_title.add_run("BẢN THUYẾT MINH QUYẾT TOÁN")
run.font.size = Pt(14)
run.font.name = 'Times New Roman'
run.font.bold = True

# === NỘI DUNG ===
def add_content_line(text, bold_label=True, indent=0):
    p = doc.add_paragraph()
    p.space_before = Pt(2)
    p.space_after = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    if ':' in text and bold_label:
        parts = text.split(':', 1)
        run1 = p.add_run(parts[0] + ':')
        run1.font.size = Pt(12)
        run1.font.name = 'Times New Roman'
        run1.font.bold = False
        if len(parts) > 1:
            run2 = p.add_run(parts[1])
            run2.font.size = Pt(12)
            run2.font.name = 'Times New Roman'
    else:
        run = p.add_run(text)
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    return p

def format_money(val):
    if val == 0: return '0'
    return f'{val:,}'.replace(',', '.')

def format_trieu(val):
    if val == 0: return '0'
    trieu = val / 1000000
    if trieu == int(trieu):
        return f'{int(trieu):,}'.replace(',', '.')
    return f'{trieu:,.2f}'.replace(',', '.')

add_content_line(f'- Tên công trình: {ten_ct}')
add_content_line(f'- MCT: {ma_ct}')
add_content_line(f'- Giá trị vốn kế hoạch: {format_money(ke_hoach)} đồng ({format_trieu(ke_hoach)} triệu đồng)')
add_content_line(f'- Thuộc kế hoạch vốn:                          Sửa chữa lớn năm {now.year}')
add_content_line(f'- Hình thức tự làm hay thuê ngoài:')
add_content_line(f'+ Tên đơn vị thi công:                          {don_vi_ql}', indent=0.5)
add_content_line(f'- Giá trị dự toán được duyệt: {format_money(gt_dt_scl)} đồng')
add_content_line(f'- Thời gian khởi công: {ngay_kc_str}')
add_content_line(f'- Thời gian hoàn thành: {ngay_ht_str}')
add_content_line(f'- Giá trị quyết toán hoàn thành: {format_money(gt_qt_scl)} đồng')

# Khối lượng công việc
add_content_line('- Khối lượng công việc chủ yếu đã tiến hành :')
if klcv:
    for line in klcv.split('\n'):
        line = line.strip()
        if line:
            p = doc.add_paragraph()
            p.space_before = Pt(1)
            p.space_after = Pt(1)
            p.paragraph_format.left_indent = Cm(0.5)
            run = p.add_run(line)
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
else:
    for _ in range(3):
        doc.add_paragraph()

# Căn cứ pháp lý
add_content_line('- Các căn cứ về chế độ để lập quyết toán:')
if can_cu_pl:
    for line in can_cu_pl.split('\n'):
        line = line.strip()
        if line:
            p = doc.add_paragraph()
            p.space_before = Pt(1)
            p.space_after = Pt(1)
            p.paragraph_format.left_indent = Cm(0.5)
            run = p.add_run(line)
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
else:
    for _ in range(3):
        doc.add_paragraph()

# Phân tích
add_content_line('- Phân tích các nhân tố tăng giảm so với dự toán được duyệt.')
add_content_line('- Đánh giá hiệu quả của công việc sửa chữa lớn (hiệu quả của việc thay thế các thiết bị so với sửa chữa các thiết bị đã hư hỏng và hiệu quả khôi phục tính năng của tài sản cố định nói chung sau khi sửa chữa:')
add_content_line('Đạt yêu cầu', bold_label=False)
add_content_line('- Các kiến nghị (nếu có)')

# Phần ký tên
doc.add_paragraph()

sign_table = doc.add_table(rows=1, cols=2)
sign_table.alignment = WD_TABLE_ALIGNMENT.CENTER

cell_noi_nhan = sign_table.cell(0, 0)
p = cell_noi_nhan.paragraphs[0]
run = p.add_run('Nơi nhận:')
run.font.size = Pt(11)
run.font.name = 'Times New Roman'
run.font.bold = True
run.font.italic = True
p2 = cell_noi_nhan.add_paragraph()
run2 = p2.add_run('- Như trên')
run2.font.size = Pt(10)
run2.font.name = 'Times New Roman'
run2.font.italic = True
p3 = cell_noi_nhan.add_paragraph()
run3 = p3.add_run('- Lưu')
run3.font.size = Pt(10)
run3.font.name = 'Times New Roman'
run3.font.italic = True

cell_gd = sign_table.cell(0, 1)
p = cell_gd.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('GIÁM ĐỐC')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
run.font.bold = True

remove_table_borders(sign_table)

# Lưu file
output_path = f"Thuyet_minh_QT_{ten_ct[:30].replace('/', '_').replace(chr(92), '_').replace(':', '_')}.docx"
doc.save(output_path)
print(f"Đã lưu file: {output_path}")
